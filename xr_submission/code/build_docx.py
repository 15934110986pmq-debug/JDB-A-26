# -*- coding: utf-8 -*-
"""md → docx 转换器（XR 第二版交付用）

按全国大学生数学建模竞赛论文格式规范（2026 年修订稿）输出：
  - 电子版第一页 = 摘要页
  - 正文不含目录
  - 附录含全部可运行源程序
  - 不出现参赛者身份/学校/赛区信息
  - 单文件 docx，≤20 MB

不依赖 pandoc。LaTeX 公式做 Unicode 替换；表格、列表、代码块、加粗、斜体保留。
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement


BASE = Path("/home/p/JDB/xr_submission")
PAPER_MD = BASE / "paper" / "A题_XR第二版.md"
OUT_DOCX = BASE / "paper" / "A题_XR第二版.docx"
CODE_DIR = BASE / "code"
SOURCE_PROGRAMS = ["program1.py", "program2.py", "program3.py", "program4.py"]


# ---------- LaTeX → Unicode 映射 ----------
MATHBF_MAP = {
    "a": "𝐚", "b": "𝐛", "c": "𝐜", "d": "𝐝", "e": "𝐞", "f": "𝐟", "g": "𝐠",
    "h": "𝐡", "i": "𝐢", "j": "𝐣", "k": "𝐤", "l": "𝐥", "m": "𝐦", "n": "𝐧",
    "o": "𝐨", "p": "𝐩", "q": "𝐪", "r": "𝐫", "s": "𝐬", "t": "𝐭", "u": "𝐮",
    "v": "𝐯", "w": "𝐰", "x": "𝐱", "y": "𝐲", "z": "𝐳",
    "X": "𝐗", "Y": "𝐘", "0": "𝟎", "1": "𝟏",
}

GREEK_NAME_TO_CHAR = {
    "tau": "τ", "sigma": "σ", "Sigma": "Σ", "rho": "ρ", "delta": "δ", "Delta": "Δ",
    "varepsilon": "ε", "epsilon": "ε", "eta": "η", "theta": "θ", "Theta": "Θ",
    "alpha": "α", "beta": "β", "gamma": "γ", "Gamma": "Γ", "lambda": "λ", "Lambda": "Λ",
    "mu": "μ", "nu": "ν", "xi": "ξ", "Xi": "Ξ", "pi": "π", "Pi": "Π", "chi": "χ",
    "phi": "φ", "Phi": "Φ", "psi": "ψ", "Psi": "Ψ", "omega": "ω", "Omega": "Ω",
    "kappa": "κ", "zeta": "ζ",
}


def _name_to_char(name: str) -> str:
    return GREEK_NAME_TO_CHAR.get(name, name)


LATEX_TOKENS = [
    # 角度符号必须最先处理，避免后续 ^\circ → ^°
    (r"\^\\circ", "°"),
    (r"\\widehat\{([^{}]+)\}", lambda m: m.group(1) + "̂"),
    (r"\\hat\{([^{}]+)\}", lambda m: m.group(1) + "̂"),
    (r"\\bar\{([^{}]+)\}", lambda m: m.group(1) + "̄"),
    (r"\\tilde\{([^{}]+)\}", lambda m: m.group(1) + "̃"),
    # 不带花括号: \widehat\tau, \bar p, \hat\Sigma 等 — 同时把 LaTeX 名转希腊字母
    (r"\\widehat\\([a-zA-Z]+)", lambda m: _name_to_char(m.group(1)) + "̂"),
    (r"\\hat\\([a-zA-Z]+)", lambda m: _name_to_char(m.group(1)) + "̂"),
    (r"\\bar\\([a-zA-Z]+)", lambda m: _name_to_char(m.group(1)) + "̄"),
    (r"\\widehat\s+([a-zA-Z])", lambda m: _name_to_char(m.group(1)) + "̂"),
    (r"\\hat\s+([a-zA-Z])", lambda m: _name_to_char(m.group(1)) + "̂"),
    (r"\\bar\s+([a-zA-Z])", lambda m: _name_to_char(m.group(1)) + "̄"),
    (r"\\mathbf\{([^{}]+)\}", lambda m: MATHBF_MAP.get(m.group(1), m.group(1))),
    (r"\\boldsymbol\{\\varepsilon\}", "𝛆"),
    (r"\\boldsymbol\{\\eta\}", "𝛈"),
    (r"\\boldsymbol\{([^{}]+)\}", lambda m: m.group(1)),
    # 不带花括号: \boldsymbol\varepsilon, \boldsymbol\eta
    (r"\\boldsymbol\\varepsilon", "𝛆"),
    (r"\\boldsymbol\\eta", "𝛈"),
    (r"\\boldsymbol\\([a-zA-Z]+)", lambda m: _name_to_char(m.group(1))),
    # 保留花括号: 后续 _/^ 正则才能正确吃到完整下/上标
    (r"\\mathrm\{([^{}]+)\}", lambda m: "{" + m.group(1) + "}"),
    (r"\\mathbb\{([^{}]+)\}", lambda m: {"E": "𝔼", "R": "ℝ", "N": "ℕ", "Z": "ℤ", "1": "𝟙"}.get(m.group(1), "{" + m.group(1) + "}")),
    (r"\\mathcal\{([^{}]+)\}", lambda m: "{" + m.group(1) + "}"),
    (r"\\text\{([^{}]+)\}", lambda m: "{" + m.group(1) + "}"),
    (r"\\boxed\{([^{}]+)\}", lambda m: "【" + m.group(1) + "】"),
    (r"\\frac\{([^{}]+)\}\{([^{}]+)\}", lambda m: "(" + m.group(1) + ")/(" + m.group(2) + ")"),
    (r"\\sqrt\{([^{}]+)\}", lambda m: "√(" + m.group(1) + ")"),
    (r"\\dot\{([^{}]+)\}", lambda m: m.group(1) + "̇"),
    (r"\\ddot\{([^{}]+)\}", lambda m: m.group(1) + "̈"),
    (r"\\vec\{([^{}]+)\}", lambda m: m.group(1) + "⃗"),
]

LATEX_SYMBOLS = {
    r"\tau": "τ", r"\sigma": "σ", r"\Sigma": "Σ", r"\rho": "ρ",
    r"\delta": "δ", r"\Delta": "Δ", r"\varepsilon": "ε", r"\eta": "η",
    r"\theta": "θ", r"\Theta": "Θ", r"\alpha": "α", r"\beta": "β", r"\gamma": "γ", r"\Gamma": "Γ",
    r"\lambda": "λ", r"\Lambda": "Λ", r"\mu": "μ", r"\pi": "π", r"\Pi": "Π", r"\chi": "χ",
    r"\phi": "φ", r"\Phi": "Φ", r"\psi": "ψ", r"\Psi": "Ψ",
    r"\omega": "ω", r"\Omega": "Ω", r"\xi": "ξ", r"\Xi": "Ξ", r"\zeta": "ζ", r"\nu": "ν", r"\kappa": "κ",
    r"\le": "≤", r"\leq": "≤", r"\ge": "≥", r"\geq": "≥",
    r"\ne": "≠", r"\neq": "≠", r"\to": "→", r"\rightarrow": "→", r"\Rightarrow": "⇒",
    r"\in": "∈", r"\notin": "∉", r"\subset": "⊂", r"\supset": "⊃",
    r"\cup": "∪", r"\cap": "∩", r"\emptyset": "∅",
    r"\sum": "∑", r"\prod": "∏", r"\int": "∫",
    r"\bigwedge": "⋀", r"\bigvee": "⋁", r"\bigcap": "⋂", r"\bigcup": "⋃",
    r"\cdot": "·", r"\cdots": "⋯", r"\ldots": "…", r"\times": "×", r"\pm": "±",
    r"\approx": "≈", r"\equiv": "≡", r"\sim": "∼", r"\propto": "∝",
    r"\infty": "∞", r"\partial": "∂", r"\nabla": "∇", r"\circ": "°", r"\degree": "°",
    r"\bmod": "mod", r"\pmod": "mod", r"\bigg": "", r"\Big": "", r"\big": "", r"\Bigg": "",
    r"\left": "", r"\right": "",
    r"\quad": "    ", r"\qquad": "        ", r"\,": " ", r"\ ": " ", r"\;": " ", r"\!": "",
    r"\top": "ᵀ", r"\dagger": "†",
    r"\min": "min", r"\max": "max", r"\sup": "sup", r"\inf": "inf",
    r"\log": "log", r"\ln": "ln", r"\exp": "exp", r"\sin": "sin", r"\cos": "cos", r"\tan": "tan",
    r"\gg": "≫", r"\ll": "≪", r"\det": "det", r"\arg": "arg",
    r"\forall": "∀", r"\exists": "∃", r"\implies": "⇒",
    r"\tfrac": "/", r"\dfrac": "/",
}

SUPERSCRIPT = {
    "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴", "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹",
    "+": "⁺", "-": "⁻", "=": "⁼", "(": "⁽", ")": "⁾",
    "a": "ᵃ", "b": "ᵇ", "c": "ᶜ", "d": "ᵈ", "e": "ᵉ", "f": "ᶠ", "g": "ᵍ",
    "h": "ʰ", "i": "ⁱ", "j": "ʲ", "k": "ᵏ", "l": "ˡ", "m": "ᵐ", "n": "ⁿ",
    "o": "ᵒ", "p": "ᵖ", "r": "ʳ", "s": "ˢ", "t": "ᵗ", "u": "ᵘ", "v": "ᵛ",
    "w": "ʷ", "x": "ˣ", "y": "ʸ", "z": "ᶻ",
    "T": "ᵀ", "A": "ᴬ", "B": "ᴮ", "D": "ᴰ", "E": "ᴱ", "G": "ᴳ",
    "*": "*",
}

SUBSCRIPT = {
    "0": "₀", "1": "₁", "2": "₂", "3": "₃", "4": "₄", "5": "₅", "6": "₆", "7": "₇", "8": "₈", "9": "₉",
    "+": "₊", "-": "₋", "=": "₌", "(": "₍", ")": "₎",
    "a": "ₐ", "e": "ₑ", "h": "ₕ", "i": "ᵢ", "j": "ⱼ", "k": "ₖ", "l": "ₗ",
    "m": "ₘ", "n": "ₙ", "o": "ₒ", "p": "ₚ", "r": "ᵣ", "s": "ₛ", "t": "ₜ",
    "u": "ᵤ", "v": "ᵥ", "x": "ₓ",
}


def _try_translate(s: str, mapping: dict) -> str | None:
    """全部字符都能映射才返回结果，否则返回 None（保留原样）。"""
    out = []
    for ch in s:
        if ch in mapping:
            out.append(mapping[ch])
        else:
            return None
    return "".join(out)


def _take_braced(text: str, start: int) -> tuple[str, int] | None:
    """从 text[start] = '{' 起读到匹配的 '}'，返回 (内部内容, 结束位置 之后)。"""
    if start >= len(text) or text[start] != "{":
        return None
    depth = 1
    i = start + 1
    while i < len(text) and depth > 0:
        c = text[i]
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0:
                return text[start + 1:i], i + 1
        i += 1
    return None


def _replace_balanced(text: str, command: str, render):
    """处理 \\command{arg}（或 \\command{arg1}{arg2}）这样的可嵌套结构。"""
    out = []
    i = 0
    needle = "\\" + command
    nl = len(needle)
    n = len(text)
    while i < n:
        if text.startswith(needle, i) and (i + nl >= n or not text[i + nl].isalpha()):
            j = i + nl
            args = []
            ok = True
            # 收集所有连续的 {arg}
            while j < n and text[j] == "{":
                got = _take_braced(text, j)
                if got is None:
                    ok = False
                    break
                arg, j = got
                args.append(arg)
                if not render.takes_more(args):
                    break
            if ok and render.accepts(args):
                out.append(render(args))
                i = j
                continue
        out.append(text[i])
        i += 1
    return "".join(out)


class _Renderer:
    def __init__(self, fn, n_args, accept=None, more=None):
        self.fn = fn
        self.n_args = n_args
        self._accept = accept or (lambda a: len(a) == n_args)
        self._more = more or (lambda a: len(a) < n_args)

    def __call__(self, args):
        return self.fn(args)

    def accepts(self, args):
        return self._accept(args)

    def takes_more(self, args):
        return self._more(args)


BALANCED_COMMANDS = [
    ("frac", _Renderer(lambda a: "(" + a[0] + ")/(" + a[1] + ")", 2)),
    ("boxed", _Renderer(lambda a: "【" + a[0] + "】", 1)),
    ("fbox", _Renderer(lambda a: "[" + a[0] + "]", 1)),
    ("dfrac", _Renderer(lambda a: "(" + a[0] + ")/(" + a[1] + ")", 2)),
    ("tfrac", _Renderer(lambda a: "(" + a[0] + ")/(" + a[1] + ")", 2)),
]


PLACEHOLDER_LBRACE = "\x01"
PLACEHOLDER_RBRACE = "\x02"


def latex_to_unicode(text: str) -> str:
    """LaTeX → Unicode。多轮迭代直到稳定，正确处理嵌套结构。"""

    # 0. 把 \{ 和 \} 暂存为占位，避免与结构化花括号混淆
    text = text.replace(r"\{", PLACEHOLDER_LBRACE).replace(r"\}", PLACEHOLDER_RBRACE)

    # 0.5. begin{cases}...end{cases} 简化为多行
    text = re.sub(r"\\begin\{cases\}(.*?)\\end\{cases\}",
                  lambda m: "{ " + m.group(1).replace(r"\\", "; ").replace("&", "  ") + " }",
                  text, flags=re.DOTALL)

    for _ in range(8):
        prev = text
        for pattern, repl in LATEX_TOKENS:
            text = re.sub(pattern, repl, text) if callable(repl) else re.sub(pattern, repl, text)
        # 平衡花括号命令
        for cmd, renderer in BALANCED_COMMANDS:
            text = _replace_balanced(text, cmd, renderer)
        if text == prev:
            break

    # 按长度降序替换，避免 \top 被 \to 抢前缀
    for sym in sorted(LATEX_SYMBOLS, key=len, reverse=True):
        text = text.replace(sym, LATEX_SYMBOLS[sym])

    # 显式上下标 (^{...}, _{...}) — 仅当全部字符可映射才转
    def repl_sup_brace(m):
        out = _try_translate(m.group(1), SUPERSCRIPT)
        return out if out is not None else "^(" + m.group(1) + ")"

    def repl_sub_brace(m):
        out = _try_translate(m.group(1), SUBSCRIPT)
        # 多字符下标若不可全部映射，按"标签"语义直接拼接（如 BICeff）
        return out if out is not None else m.group(1)

    text = re.sub(r"\^\{([^{}]+)\}", repl_sup_brace, text)
    text = re.sub(r"_\{([^{}]+)\}", repl_sub_brace, text)

    # 单字符上下标
    def repl_sup_char(m):
        c = m.group(1)
        return SUPERSCRIPT[c] if c in SUPERSCRIPT else "^" + c

    def repl_sub_char(m):
        c = m.group(1)
        return SUBSCRIPT[c] if c in SUBSCRIPT else "_" + c

    text = re.sub(r"\^(\w)", repl_sup_char, text)
    text = re.sub(r"_(\w)", repl_sub_char, text)

    text = text.replace("\\\\", "\n").replace("\\&", "&").replace(r"\%", "%").replace(r"\$", "$").replace(r"\#", "#")
    # 残留的 \xxx 控制序列，移除反斜杠
    text = re.sub(r"\\([a-zA-Z]+)", r"\1", text)

    text = text.replace("{", "").replace("}", "")
    # 还原 \{ \} 占位
    text = text.replace(PLACEHOLDER_LBRACE, "{").replace(PLACEHOLDER_RBRACE, "}")
    text = text.replace("$", "")
    return text


# ---------- inline 解析（**bold** *italic* `code` $math$） ----------

INLINE_PATTERN = re.compile(
    r"(\$[^$\n]+\$|\*\*[^*]+\*\*|\*[^*\n]+\*|`[^`\n]+`)"
)


def add_inline_runs(paragraph, text: str, base_bold=False, base_italic=False):
    if not text:
        return
    parts = INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("$") and part.endswith("$"):
            run = paragraph.add_run(latex_to_unicode(part[1:-1]))
            run.italic = True
            run.font.name = "Cambria Math"
        elif part.startswith("**") and part.endswith("**"):
            # 粗体内部递归，处理内嵌 $...$ / *...* / `...`
            add_inline_runs(paragraph, part[2:-2], base_bold=True, base_italic=base_italic)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            add_inline_runs(paragraph, part[1:-1], base_bold=base_bold, base_italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.bold = base_bold
            run.italic = base_italic
        else:
            run = paragraph.add_run(part)
            run.bold = base_bold
            run.italic = base_italic


# ---------- markdown 行扫描 ----------

def parse_md(md_text: str):
    """Yield 结构化 token: ('h', level, text) / ('p', text) / ('disp', tex) /
       ('table', rows) / ('ul', items) / ('ol', items) / ('code', text) / ('hr',) / ('blank',)
    """
    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        s = line.rstrip()

        if not s.strip():
            yield ("blank",)
            i += 1
            continue

        # 水平线 (---) → 跳过，我们用空行分隔
        if re.match(r"^-{3,}\s*$", s):
            yield ("blank",)
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            yield ("h", len(m.group(1)), m.group(2))
            i += 1
            continue

        # 显示公式 $$...$$（单行或多行）
        if s.strip().startswith("$$"):
            stripped = s.strip()
            # 单行：$$...$$
            if len(stripped) >= 4 and stripped.endswith("$$") and stripped.count("$$") >= 2:
                yield ("disp", stripped[2:-2])
                i += 1
                continue
            # 多行：开始行只含 $$ 或 $$+起始内容
            buf = [stripped[2:]] if len(stripped) > 2 else []
            i += 1
            while i < n and "$$" not in lines[i]:
                buf.append(lines[i])
                i += 1
            if i < n:
                tail = lines[i].rstrip()
                if tail.endswith("$$"):
                    buf.append(tail[:-2])
                else:
                    buf.append(tail.replace("$$", ""))
                i += 1
            yield ("disp", "\n".join(b for b in buf if b))
            continue

        # 代码块 ```
        if s.strip().startswith("```"):
            buf = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            yield ("code", "\n".join(buf))
            continue

        # 表格 |...|
        if s.lstrip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1]):
            rows = []
            header = [c.strip() for c in s.strip().strip("|").split("|")]
            rows.append(header)
            i += 2
            while i < n and lines[i].lstrip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            yield ("table", rows)
            continue

        # 无序列表
        if re.match(r"^[\-\*]\s+", s):
            items = []
            while i < n and re.match(r"^[\-\*]\s+", lines[i]):
                items.append(re.sub(r"^[\-\*]\s+", "", lines[i].rstrip()))
                i += 1
            yield ("ul", items)
            continue

        # 有序列表
        if re.match(r"^\d+\.\s+", s):
            items = []
            while i < n and re.match(r"^\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].rstrip()))
                i += 1
            yield ("ol", items)
            continue

        # 段落（合并连续非空行）
        buf = [s]
        i += 1
        while i < n and lines[i].strip() and not (
            lines[i].lstrip().startswith("#")
            or lines[i].lstrip().startswith("|")
            or lines[i].strip().startswith("$$")
            or lines[i].strip().startswith("```")
            or re.match(r"^[\-\*]\s+", lines[i])
            or re.match(r"^\d+\.\s+", lines[i])
            or re.match(r"^-{3,}\s*$", lines[i])
        ):
            buf.append(lines[i].rstrip())
            i += 1
        yield ("p", " ".join(buf))


# ---------- 写 docx ----------

def setup_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts")) or OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    if rpr.find(qn("w:rFonts")) is None:
        rpr.insert(0, rfonts)


def add_heading(doc: Document, level: int, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    if level == 1:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 18
    elif level == 2:
        size = 16
    elif level == 3:
        size = 14
    else:
        size = 12
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "黑体"
    rpr = run.element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "黑体")
    rpr.append(rfonts)


def add_paragraph(doc: Document, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5
    add_inline_runs(para, text)


def add_display_math(doc: Document, tex: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(latex_to_unicode(tex))
    run.italic = True
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    n_col = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_col)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(rows):
        for c in range(n_col):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            text = row[c] if c < len(row) else ""
            for para in cell.paragraphs:
                para._element.getparent().remove(para._element)
            para = cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(para, text, base_bold=(r == 0))


def add_list(doc: Document, items: list[str], ordered: bool):
    style = "List Number" if ordered else "List Bullet"
    for it in items:
        para = doc.add_paragraph(style=style)
        add_inline_runs(para, it)


def add_code_block(doc: Document, text: str):
    for line in text.splitlines() or [""]:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def add_image_if_exists(doc: Document, path: Path, caption: str = ""):
    if not path.exists():
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(path), width=Cm(13.5))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True
            r.font.size = Pt(10)


# ---------- 主流程 ----------

def render_paper(doc: Document):
    md = PAPER_MD.read_text(encoding="utf-8")
    tokens = list(parse_md(md))

    in_section_5_4_10 = False
    figures_inserted = set()

    for tok in tokens:
        kind = tok[0]
        if kind == "h":
            level, text = tok[1], tok[2]
            add_heading(doc, level, text)
            # 在对应章节末尾插入图
            if level == 4 and "5.4.10" in text:
                in_section_5_4_10 = True
        elif kind == "p":
            add_paragraph(doc, tok[1])
        elif kind == "disp":
            add_display_math(doc, tok[1])
        elif kind == "table":
            add_table(doc, tok[1])
            doc.add_paragraph()
        elif kind == "ul":
            add_list(doc, tok[1], ordered=False)
        elif kind == "ol":
            add_list(doc, tok[1], ordered=True)
        elif kind == "code":
            add_code_block(doc, tok[1])
        elif kind == "blank":
            pass

    # 插入关键诊断图（章节末尾统一附）
    fig_dir = BASE / "figures"
    figs = [
        ("q1_time_shift_objective.png", "图 5.1  附件 1 时间偏差目标函数 J(τ) 曲线"),
        ("q1_aligned_trajectory.png", "图 5.2  附件 1 校正后两路轨迹叠加"),
        ("q2_profile_objective.png", "图 5.3  附件 2 剖面目标函数 Jₑ(τ) 曲线"),
        ("q2_residual_ecdf.png", "图 5.4  附件 2 残差 ECDF（剥离偏差前后）"),
        ("q2_covariance_comparison.png", "图 5.5  附件 2 三种轨迹的 95% 椭圆面积比较"),
        ("q2_fused_10hz_trajectory.png", "图 5.6  附件 2 10Hz 融合轨迹"),
        ("q3_time_objective.png", "图 5.7  附件 3 时间偏差目标函数曲线"),
        ("q3_bias_bootstrap_ci.png", "图 5.8  附件 3 Bootstrap 候选偏差 95%CI"),
        ("q3_hac_sensitivity.png", "图 5.9  附件 3 HAC lag-scan p 值曲线"),
        ("q3_residual_ecdf.png", "图 5.10  附件 3 残差 ECDF"),
        ("q3_extended_trajectory_for_q4.png", "图 5.11  第三问扩展 10Hz 轨迹（供第四问使用）"),
        ("q4_trajectory_targets_second.png", "图 5.12  第四问轨迹与目标点叠加图"),
        ("q4_photo_counts_second.png", "图 5.13  第四问拍照覆盖统计"),
    ]
    add_heading(doc, 2, "图件汇编")
    para = doc.add_paragraph()
    add_inline_runs(para, "下列图件按章节出现顺序依次列出，与正文 §5.1–§5.4 对应。")
    for fn, cap in figs:
        add_image_if_exists(doc, fig_dir / fn, cap)


def append_source_code(doc: Document):
    """规范第五条：附录必须含全部可运行源程序。"""
    add_heading(doc, 2, "8.7 程序源代码（可运行）")
    para = doc.add_paragraph()
    add_inline_runs(para, "以下源代码与提交目录 code/ 内文件一字不差，直接 python <文件名> 即可运行。")

    for fn in SOURCE_PROGRAMS:
        path = CODE_DIR / fn
        if not path.exists():
            continue
        add_heading(doc, 3, fn)
        text = path.read_text(encoding="utf-8")
        add_code_block(doc, text)


def main():
    doc = Document()
    setup_document(doc)
    render_paper(doc)
    append_source_code(doc)
    doc.save(str(OUT_DOCX))
    size_mb = OUT_DOCX.stat().st_size / (1024 * 1024)
    print(f"OK: {OUT_DOCX}  ({size_mb:.2f} MB)")


if __name__ == "__main__":
    main()
